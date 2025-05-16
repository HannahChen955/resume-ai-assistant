#!/usr/bin/env python3

import os
import re
import json
import uuid
import fitz
import subprocess
import tempfile
from shutil import which
from docx import Document
from tqdm import tqdm
from paddleocr import PaddleOCR
from dotenv import load_dotenv
from openai import OpenAI

# ✅ 初始化
load_dotenv()

INPUT_DIR = os.getenv("RESUME_DIR", "data/resumes")
OUTPUT_DIR = os.getenv("EXTRACTED_DIR", "data/resumes_extract_enhanced")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
EXTRACTION_MODEL = os.getenv("LLM_FIELD_EXTRACT_MODEL", "gpt-3.5-turbo")
BLACKLIST = ["个人简历", "猎聘", "BOSS直聘", "客户名称", "项目名称", "original", "standard"]

openai_client = OpenAI(api_key=OPENAI_API_KEY)
ocr_model = PaddleOCR(use_textline_orientation=True, lang='ch')

# ========= PDF OCR (PaddleOCR) =========
def extract_via_ocr(file_path):
    try:
        print(f"📸 PaddleOCR 识别中：{os.path.basename(file_path)}")
        result = ocr_model.ocr(file_path, cls=True)
        lines = [line[1][0] for line in result[0] if line[1][0].strip()]
        text = "\n".join(lines)
        if len(text.strip()) < 30:
            print(f"⚠️ OCR 提取过少: {file_path}")
        return text if len(text.strip()) >= 30 else None
    except Exception as e:
        print(f"⚠️ OCR 失败: {e}")
        return None

# ========= PDF 正文提取 =========
def extract_pdf_text(file_path):
    try:
        doc = fitz.open(file_path)
        lines = []
        for page in doc:
            blocks = page.get_text("dict")["blocks"]
            for b in blocks:
                if "lines" in b:
                    for line in b["lines"]:
                        span_text = " ".join([span["text"] for span in line["spans"] if span.get("size", 0) > 8])
                        if span_text.strip():
                            lines.append(span_text.strip())
        lines = [l for l in lines if len(l.strip()) > 5]
        result = "\n".join(lines)
        if len(result.strip()) < 300:
            print(f"📉 PDF文本过少，强制OCR: {os.path.basename(file_path)}")
            return extract_via_ocr(file_path)
        return result
    except Exception as e:
        print(f"❌ PDF读取失败：{file_path} | {e}")
        return extract_via_ocr(file_path)

# ========= DOC 提取增强 =========
def extract_doc_text(file_path):
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".txt") as temp_output:
            subprocess.run(["unoconv", "-f", "txt", "-o", temp_output.name, file_path], check=True)
            with open(temp_output.name, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            os.remove(temp_output.name)

            if len(text.strip()) > 100:
                return text
            else:
                print(f"📉 DOC文本过少，转PDF后OCR兜底: {file_path}")
                pdf_path = file_path + ".pdf"
                subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(file_path), file_path], check=True)
                if os.path.exists(pdf_path):
                    ocr_text = extract_via_ocr(pdf_path)
                    os.remove(pdf_path)
                    return ocr_text
                return None
    except Exception as e:
        print(f"❌ DOC读取失败：{file_path} | {e}")
        return None

# ========= DOCX 提取 =========
def extract_docx_text(file_path):
    try:
        doc = Document(file_path)
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    tables.append(row_text)
        content = "\n".join(paras + tables)
        return content if len(content.strip()) > 10 else None
    except Exception as e:
        print(f"❌ DOCX读取失败：{file_path} | {e}")
        return None

# ========= 字段提取 =========
def extract_fields(text: str) -> dict:
    prompt = f"""请从以下文本中识别以下字段（若缺失请写 null）：
- 姓名
- 应聘职位
- 手机号
- 邮箱
仅返回标准 JSON。
文本：
\"\"\"{text[:600]}\"\"\""""
    try:
        response = openai_client.chat.completions.create(
            model=EXTRACTION_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=300
        )
        content = response.choices[0].message.content.strip()
        content = re.sub(r"^```(json)?|```$", "", content)
        return json.loads(content)
    except Exception as e:
        print(f"⚠️ 字段提取失败: {e}")
        return {}

# ========= 结构分类提取 =========
def classify_resume_sections(text: str) -> list:
    prompt = f"""请根据以下简历内容，将其分类为以下模块：
- 岗位经历
- 项目经历
- 教育背景
返回JSON数组格式。
简历内容：
\"\"\"{text[:3000]}\"\"\""""
    try:
        response = openai_client.chat.completions.create(
            model=EXTRACTION_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=1000
        )
        content = response.choices[0].message.content.strip()
        content = re.sub(r"^```(json)?|```$", "", content)
        return json.loads(content)
    except Exception as e:
        print(f"⚠️ 结构识别失败: {e}")
        return []

# ========= 清洗增强 =========
def enhance_text(text):
    text = re.sub(r'\b[a-zA-Z0-9]{20,}\b', '', text)
    text = re.sub(r'(简历来自|BOSS直聘|猎聘|前程无忧)[^\n]*', '', text)
    text = re.sub(r'(~{2,}|-{2,}|={2,}|\+{2,})', '', text)
    text = re.sub(r'\s{2,}', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

# ========= 姓名 fallback =========
def extract_name_fallback(filename, text):
    base = os.path.splitext(os.path.basename(filename))[0]
    name_candidates = re.findall(r'[\u4e00-\u9fa5]{2,4}', base)
    for name in name_candidates:
        if name not in BLACKLIST:
            return name
    head_text = text[:50]
    name_candidates = re.findall(r'[\u4e00-\u9fa5]{2,4}', head_text)
    for name in name_candidates:
        if name not in BLACKLIST:
            return name
    return "未知姓名"

# ========= 主流程 =========
def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    files = [f for f in os.listdir(INPUT_DIR) if f.lower().endswith(('.pdf', '.doc', '.docx'))]
    success, fail = 0, 0

    for filename in tqdm(files, desc="📄 正在提取简历"):
        try:
            path = os.path.join(INPUT_DIR, filename)
            ext = filename.lower().split('.')[-1]

            if ext == "pdf":
                text = extract_pdf_text(path)
            elif ext == "docx":
                text = extract_docx_text(path)
            elif ext == "doc":
                text = extract_doc_text(path)
            else:
                text = None

            if not text or len(text.strip()) < 10:
                text = "[文件读取失败或内容为空]"

            text = enhance_text(text)
            fields = extract_fields(text)
            sections = classify_resume_sections(text)

            if not fields.get("姓名"):
                fields["姓名"] = extract_name_fallback(filename, text)

            output_name = f"{fields['姓名']}_{fields.get('应聘职位', '未知职位')}_{str(uuid.uuid4())[:8]}.txt"
            output_path = os.path.join(OUTPUT_DIR, output_name)

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(text)
                f.write("\n\n=== 结构识别结果 ===\n")
                f.write(json.dumps(sections, ensure_ascii=False, indent=2))
            print(f"✅ 输出完成：{output_name}")

            success += 1

        except Exception as e:
            print(f"❌ 处理失败：{filename} | {e}")
            fail += 1

    print(f"\n🎯 总数：{len(files)} | 成功：{success} | 失败：{fail}")

if __name__ == "__main__":
    main()
