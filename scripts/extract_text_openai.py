#!/usr/bin/env python3

import os
import re
import json
import uuid
import fitz
import subprocess
import tempfile
from docx import Document
from tqdm import tqdm
from pdf2image import convert_from_path
import pytesseract
from dotenv import load_dotenv
from openai import OpenAI
from pathlib import Path
from dotenv import dotenv_values

# ✅ 初始化
load_dotenv()

root_dir = Path(__file__).resolve().parent.parent
config = dotenv_values(root_dir / ".env")

api_key = config.get("OPENAI_API_KEY")
embedding_model = config.get("OPENAI_EMBEDDING_MODEL")
llm_model = config.get("OPENAI_LLM_FIELD_EXTRACT_MODEL")
ie_model = None

INPUT_DIR = os.getenv("RESUME_DIR", "data/resumes")
OUTPUT_DIR = os.getenv("EXTRACTED_DIR", "data/resumes_extract_enhanced")
BLACKLIST = [
    "个人简历", "简历", "求职", "BOSS直聘", "猎聘", "客户名称", "项目名称",
    "original", "standard", "的", "doc", "pdf", "docx", "附件"
]

openai_client = OpenAI(api_key=api_key)

# ========= OCR 兜底 =========
def extract_via_ocr(file_path):
    try:
        print(f"📸 OCR 识别中：{os.path.basename(file_path)}")
        images = convert_from_path(file_path, dpi=300)
        results = [pytesseract.image_to_string(img, lang='chi_sim', config='--psm 6 --oem 3') for img in images]
        text_result = "\n".join([t.strip() for t in results if t.strip()])
        return text_result if len(text_result.strip()) >= 30 else None
    except Exception as e:
        print(f"⚠️ OCR 失败: {e}")
        return None

# ========= 文本提取 =========
def extract_text(file_path):
    ext = file_path.lower().split('.')[-1]
    if ext == "pdf":
        return extract_pdf_text(file_path)
    elif ext == "docx":
        text = extract_docx_text(file_path)
        if not text or len(text.strip()) < 30:
            pdf_path = convert_to_pdf(file_path)
            return extract_via_ocr(pdf_path)
        return text
    elif ext == "doc":
        return extract_doc_text(file_path)
    return None

def convert_to_pdf(docx_path):
    pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
    try:
        subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(docx_path), docx_path], check=True)
        return pdf_path if os.path.exists(pdf_path) else docx_path
    except:
        return docx_path

def extract_pdf_text(file_path):
    try:
        doc = fitz.open(file_path)
        lines = []
        for page in doc:
            blocks = page.get_text("dict")["blocks"]
            for b in blocks:
                if "lines" in b:
                    for line in b["lines"]:
                        span_text = " ".join([span["text"] for span in line["spans"] if span.get("size", 0) > 8])
                        if span_text.strip():
                            lines.append(span_text.strip())
        result = "\n".join(lines)
        if len(result.strip()) < 300:
            return extract_via_ocr(file_path)
        return result
    except:
        return extract_via_ocr(file_path)

def extract_doc_text(file_path):
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".txt") as temp_output:
            subprocess.run(["unoconv", "-f", "txt", "-o", temp_output.name, file_path], check=True)
            with open(temp_output.name, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            os.remove(temp_output.name)
            if len(text.strip()) > 100:
                return text
            pdf_path = convert_to_pdf(file_path)
            return extract_via_ocr(pdf_path)
    except:
        return None

def extract_docx_text(file_path):
    try:
        doc = Document(file_path)
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    tables.append(row_text)
        return "\n".join(paras + tables)
    except:
        return None

# ========= 清洗增强 =========
def enhance_text(text):
    text = re.sub(r'\b[a-zA-Z0-9]{20,}\b', '', text)
    text = re.sub(r'(简历来自|BOSS直聘|猎聘|前程无忧)[^\n]*', '', text)
    text = re.sub(r'(~{2,}|-{2,}|={2,}|\+{2,})', '', text)
    text = re.sub(r'\s{2,}', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

# ========= 字段提取 =========
def extract_fields(text):
    prompt = f"""
你是一位信息抽取专家，请从以下中文简历文本中提取出候选人的基础信息，严格按照如下字段输出 JSON（无则写 null）：

- 姓名：优先从文本前100字符中查找姓名，避免错误提取职位名或格式说明。
- 应聘职位：如"销售总监""运营经理"等角色职称。
- 手机号：中国大陆手机号，11位纯数字。
- 邮箱：包含 @ 字符，常见邮箱格式。

请你只输出标准 JSON 格式，不添加注释或解释。

示例：
{{
  "姓名": "张三",
  "应聘职位": "产品经理",
  "手机号": "13812345678",
  "邮箱": "zhangsan@example.com"
}}

以下是简历正文：
{text[:3000]}
"""
    try:
        response = openai_client.chat.completions.create(
            model=llm_model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=300
        )
        content = response.choices[0].message.content.strip()
        content = re.sub(r"^```(json)?|```$", "", content)
        return json.loads(content)
    except Exception as e:
        print(f"⚠️ 字段提取失败: {e}")
        return {}

# ========= 模块结构分类提取 =========
def classify_resume_sections(text):
    prompt = f"""
你是一位结构分析助手，请对以下简历内容进行分类整理，将其拆分为如下模块：

- 岗位经历：具体工作单位、时间、岗位及职责。
- 项目经历：参与的项目内容、成果、职责。
- 教育背景：学校、学历、专业、时间。
- 技能亮点：技能名称、掌握程度、证书等。
- 自我评价：自述型内容，包括优劣势、性格等。

要求：

- 以 JSON 数组格式返回，每个元素为一个模块。
- 每个模块包含字段："模块" 和 "内容"。
- 不包含 markdown、解释说明、额外标签。
- 模块内容尽量提炼为简洁段落，如有编号/表格可合并。

示例输出：
[
  {{
    "模块": "岗位经历",
    "内容": "2020年-2023年在某科技公司担任产品经理，负责电商平台设计。"
  }},
  {{
    "模块": "教育背景",
    "内容": "2014年-2018年，复旦大学，本科，计算机科学与技术。"
  }}
]

以下是简历内容：
{text[:3000]}
"""
    try:
        response = openai_client.chat.completions.create(
            model=llm_model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=1000
        )
        content = response.choices[0].message.content.strip()
        content = re.sub(r"^```(json)?|```$", "", content)
        return json.loads(content)
    except Exception as e:
        print(f"⚠️ 结构识别失败: {e}")
        return []

# ========= 文件名清洗 =========
def sanitize_filename_part(value: str) -> str:
    return re.sub(r'[\\/:*?"<>|]', '_', value or "null")

# ========= 姓名 fallback =========
def extract_name_fallback(filename, text):
    base = os.path.splitext(os.path.basename(filename))[0]
    first_part = base.split('_')[0]  # 通常是姓名部分
    name_candidates = re.findall(r'[\u4e00-\u9fa5]{2,4}', first_part)
    for name in name_candidates:
        if name not in BLACKLIST:
            # 清洗掉尾部的"的"或"简历"等非人名词
            name = re.sub(r'(的|简历)$', '', name)
            return name
    head_text = text[:50]
    name_candidates = re.findall(r'[\u4e00-\u9fa5]{2,4}', head_text)
    for name in name_candidates:
        if name not in BLACKLIST:
            name = re.sub(r'(的|简历)$', '', name)
            return name
    return "未知姓名"

# ========= 主流程 =========
def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    files = [f for f in os.listdir(INPUT_DIR) if f.lower().endswith(('.pdf', '.doc', '.docx'))]
    success, fail = 0, 0

    for filename in tqdm(files, desc="📄 正在提取简历"):
        try:
            path = os.path.join(INPUT_DIR, filename)
            text = extract_text(path)
            if not text or len(text.strip()) < 10:
                text = "[文件读取失败或内容为空]"

            text = enhance_text(text)
            fields = extract_fields(text)

            # fallback: 姓名不能为空
            if not fields.get("姓名") or fields["姓名"] == "null":
                fields["姓名"] = extract_name_fallback(filename, text)

            if not fields.get("应聘职位"):
                fields["应聘职位"] = "未知职位"

            name = re.findall(r'[\u4e00-\u9fa5]{2,4}', filename)[0] if re.findall(r'[\u4e00-\u9fa5]{2,4}', filename) else "未知姓名"
            position = fields.get("应聘职位", "未知职位")
            name_clean = re.sub(r'[\\/:*?"<>|]', '_', name)
            position_clean = re.sub(r'[\\/:*?"<>|]', '_', position)
            output_name = f"{name_clean}_{position_clean}_{str(uuid.uuid4())[:8]}.txt"
            output_path = os.path.join(OUTPUT_DIR, output_name)

            with open(output_path, "w", encoding="utf-8") as f:
                f.write("=== 字段提取结果 ===\n")
                f.write(json.dumps(fields, ensure_ascii=False, indent=2))
                f.write("\n\n=== 原始简历文本 ===\n")
                f.write(text)

            print(f"✅ 输出完成：{output_name}")
            success += 1

        except Exception as e:
            print(f"❌ 处理失败：{filename} | {e}")
            with open("failures.log", "a") as log:
                log.write(f"处理失败: {filename} | {e}\n")
            fail += 1

    print(f"\n🎯 总数：{len(files)} | 成功：{success} | 失败：{fail}")

if __name__ == "__main__":
    main()